import csv
from collections import OrderedDict
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
REPORT_DATE = "2026-08-13"
SOURCE = ROOT / "docs" / f"controle-meervoudige-batterijkoppelingen-{REPORT_DATE}.csv"
OUTPUT = ROOT / "docs" / f"controle-meervoudige-batterijkoppelingen-{REPORT_DATE}.docx"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
MUTED = RGBColor(89, 89, 89)
HEADER_FILL = "E8EEF5"
ALT_FILL = "F7F9FB"
WHITE = "FFFFFF"
TABLE_WIDTH_DXA = 9360
COLUMN_WIDTHS_DXA = [520, 1950, 1550, 3060, 2280]


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(TABLE_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in COLUMN_WIDTHS_DXA:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            set_cell_width(cell, COLUMN_WIDTHS_DXA[index])
            set_cell_margins(cell)


def set_run_font(run, size, bold=False, color=None, italic=False):
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def set_paragraph_spacing(paragraph, before=0, after=0, line=1.0):
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line


def add_page_field(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, separate, text, end])
    set_run_font(run, 9, color=MUTED)


def configure_styles(document):
    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ):
        style = document.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def group_rows():
    with SOURCE.open("r", encoding="utf-8-sig", newline="") as handle:
        rows = list(csv.DictReader(handle, delimiter=";"))

    groups = OrderedDict()
    for row in rows:
        groups.setdefault(row["groep"], []).append(row)
    return groups


def add_table_header(table):
    headers = ["Nr.", "Klant", "Voertuig", "Gekoppelde kaarten", "Oordeel / opmerkingen Age"]
    row = table.rows[0]
    tr_pr = row._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    tr_pr.append(repeat)

    for index, (cell, label) in enumerate(zip(row.cells, headers)):
        set_cell_shading(cell, HEADER_FILL)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if index in (0, 2) else WD_ALIGN_PARAGRAPH.LEFT
        set_paragraph_spacing(paragraph, after=0, line=1.0)
        set_run_font(paragraph.add_run(label), 9, bold=True, color=DARK_BLUE)


def add_group_row(table, group_number, rows):
    row = table.add_row()
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)
    if int(group_number) % 2 == 0:
        for cell in row.cells:
            set_cell_shading(cell, ALT_FILL)

    for cell in row.cells:
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    number_p = row.cells[0].paragraphs[0]
    number_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(number_p, after=0)
    set_run_font(number_p.add_run(group_number), 9.5, bold=True, color=DARK_BLUE)

    customer_p = row.cells[1].paragraphs[0]
    set_paragraph_spacing(customer_p, after=0, line=1.0)
    set_run_font(customer_p.add_run(rows[0]["klant"]), 9)

    vehicle_p = row.cells[2].paragraphs[0]
    vehicle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(vehicle_p, after=1, line=1.0)
    set_run_font(vehicle_p.add_run(rows[0]["voertuig_machine_nummer"]), 9, bold=True)
    type_p = row.cells[2].add_paragraph()
    type_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(type_p, after=0, line=1.0)
    set_run_font(type_p.add_run(rows[0]["voertuig_merk_model"]), 8, color=MUTED)

    cards_cell = row.cells[3]
    cards_cell.paragraphs[0]._element.getparent().remove(cards_cell.paragraphs[0]._element)
    for index, item in enumerate(rows, start=1):
        paragraph = cards_cell.add_paragraph()
        set_paragraph_spacing(paragraph, after=2 if index < len(rows) else 0, line=1.0)
        archived_label = " [ARCHIEF]" if item["batterij_lader_gearchiveerd"] == "ja" else ""
        title = item["batterij_merk_type"] or item["batterij_lader_machine_nummer"]
        set_run_font(paragraph.add_run(f"{index}. {title}{archived_label}"), 8.5, bold=True)
        charger = item["lader_merk_type"]
        history = item["keuringshistorie_batterij_lader"]
        details = " | ".join(part for part in (charger, history) if part)
        if details:
            detail_p = cards_cell.add_paragraph()
            set_paragraph_spacing(detail_p, after=3 if index < len(rows) else 0, line=1.0)
            set_run_font(detail_p.add_run(details), 7.5, color=MUTED)

    review_cell = row.cells[4]
    review_p = review_cell.paragraphs[0]
    set_paragraph_spacing(review_p, after=4, line=1.0)
    set_run_font(review_p.add_run("[ ] klopt   [ ] corrigeren"), 8.5)
    for _ in range(3):
        line_p = review_cell.add_paragraph("________________________________")
        set_paragraph_spacing(line_p, after=3, line=1.0)
        for run in line_p.runs:
            set_run_font(run, 8, color=MUTED)


def build_document():
    document = Document()
    configure_styles(document)
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    header_p = section.header.paragraphs[0]
    set_paragraph_spacing(header_p, after=0, line=1.0)
    set_run_font(header_p.add_run("Controle batterij-/laderkoppelingen"), 9, bold=True, color=MUTED)

    footer_p = section.footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_paragraph_spacing(footer_p, after=0, line=1.0)
    set_run_font(footer_p.add_run("Pagina "), 9, color=MUTED)
    add_page_field(footer_p)

    title = document.add_paragraph()
    set_paragraph_spacing(title, before=0, after=4, line=1.0)
    set_run_font(title.add_run("Controle meervoudige batterij-/laderkoppelingen"), 23, bold=True, color=DARK_BLUE)

    subtitle = document.add_paragraph()
    set_paragraph_spacing(subtitle, after=10, line=1.0)
    set_run_font(subtitle.add_run("Productiestand 13 augustus 2026"), 11, color=MUTED)

    intro = document.add_paragraph()
    set_paragraph_spacing(intro, after=10, line=1.25)
    set_run_font(
        intro.add_run(
            "In deze lijst staan alleen voertuigen met meer dan een gekoppelde batterij-/laderkaart. "
            "Controleer per groep of alle kaarten behouden moeten blijven of dat een kaart dubbel of verkeerd gekoppeld is."
        ),
        10.5,
    )

    totals = document.add_paragraph()
    set_paragraph_spacing(totals, after=12, line=1.0)
    set_run_font(totals.add_run("19 voertuigen | 42 gekoppelde kaarten"), 10.5, bold=True, color=DARK_BLUE)

    groups = group_rows()
    table = document.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    add_table_header(table)
    for group_number, rows in groups.items():
        add_group_row(table, group_number, rows)
    set_table_geometry(table)

    document.core_properties.title = "Controle meervoudige batterij-/laderkoppelingen"
    document.core_properties.subject = "Printlijst voor controle door Age"
    document.core_properties.author = "Heftrucks Friesland"
    document.core_properties.keywords = "batterij, lader, koppelingen, controle"
    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_document()
